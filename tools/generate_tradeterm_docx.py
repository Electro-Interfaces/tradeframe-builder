# -*- coding: utf-8 -*-
"""
Скрипт для генерации Word документа TradeTerm из Markdown
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Пути
PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMAGES_DIR = os.path.join(PROJECT_DIR, "TerminalTCO")
OUTPUT_FILE = os.path.join(PROJECT_DIR, "docs", "user-guide", "TradeTerm_Guide.docx")

def add_heading(doc, text, level=1):
    """Добавить заголовок"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    """Добавить параграф"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_image(doc, image_path, width=Inches(5.5)):
    """Добавить изображение"""
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=width)
        # Центрировать последний параграф (картинку)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"[Изображение не найдено: {image_path}]")

def add_table(doc, headers, rows):
    """Добавить таблицу"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Заголовки
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Жирный шрифт для заголовков
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Данные
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = cell_text

    return table

def create_document():
    """Создать Word документ"""
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # ===== ТИТУЛЬНАЯ СТРАНИЦА =====
    title = doc.add_heading('TradeTerm', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Терминал самообслуживания')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)

    doc.add_paragraph()
    subtitle2 = doc.add_paragraph('Руководство для менеджеров и операторов АЗС')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(14)

    doc.add_page_break()

    # ===== СОДЕРЖАНИЕ =====
    add_heading(doc, 'Содержание', 1)
    contents = [
        '1. Что такое TradeTerm?',
        '2. Экосистема TradeSuite',
        '3. Варианты применения TradeTerm',
        '4. Инструкция по использованию терминала',
        '5. Экраны терминала'
    ]
    for item in contents:
        doc.add_paragraph(item)

    doc.add_page_break()

    # ===== РАЗДЕЛ 1: ЧТО ТАКОЕ TRADETERM =====
    add_heading(doc, '1. Что такое TradeTerm?', 1)

    doc.add_paragraph(
        'TradeTerm (терминал самообслуживания) — это программно-аппаратный комплекс, '
        'который позволяет клиентам самостоятельно оплачивать топливо без помощи оператора. '
        'TradeTerm может работать как основная система управления заправочной станцией '
        'или как дополнительное оборудование.'
    )

    add_heading(doc, 'Преимущества для бизнеса', 2)
    add_table(doc,
        ['Преимущество', 'Описание'],
        [
            ['Снижение очередей', 'Клиенты не ждут оператора'],
            ['Работа 24/7', 'Терминал работает круглосуточно без персонала'],
            ['Экономия на персонале', 'Меньше операторов или полностью без них'],
            ['Контроль в реальном времени', 'Все операции видны в TradeFrame'],
            ['Удалённое управление', 'Контроль резервуаров через интернет'],
        ]
    )

    doc.add_paragraph()
    add_heading(doc, 'Способы оплаты на терминале', 2)

    payment_methods = [
        ('Наличные', 'оплата купюрами через купюроприёмник'),
        ('Банковская карта', 'Visa, MasterCard, МИР и другие'),
        ('Топливная карта', 'корпоративные и дисконтные карты'),
        ('Онлайн-заказ', 'оплата через мобильное приложение'),
    ]
    for method, desc in payment_methods:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(method).bold = True
        p.add_run(f' — {desc}')

    doc.add_page_break()

    # ===== РАЗДЕЛ 2: ЭКОСИСТЕМА TRADESUITE =====
    add_heading(doc, '2. Экосистема TradeSuite', 1)

    doc.add_paragraph(
        'TradeSuite — зонтичный бренд, объединяющий комплекс решений для топливного бизнеса. '
        'TradeTerm является одним из продуктов этой экосистемы.'
    )

    add_heading(doc, '2.1 Продукты экосистемы', 2)
    add_table(doc,
        ['Продукт', 'Назначение'],
        [
            ['TradeFrame Builder', 'Платформа управления станциями'],
            ['TradePOS', 'Классические АЗС с оператором'],
            ['TradeGas', 'АГЗС для СУГ (пропан-бутан) с оператором'],
            ['TradeTerm', 'Терминалы самообслуживания / АКАЗС'],
            ['TradeCorp', 'Корпоративные топливные карты (B2B)'],
            ['TradeBonus', 'Программа лояльности (B2C)'],
            ['TradeGate', 'Интеграция с агрегаторами и онлайн-заказы'],
        ]
    )

    doc.add_page_break()

    # ===== РАЗДЕЛ 3: ВАРИАНТЫ ПРИМЕНЕНИЯ =====
    add_heading(doc, '3. Варианты применения TradeTerm', 1)

    doc.add_paragraph('TradeTerm может использоваться в двух основных режимах:')

    add_heading(doc, '3.1 АКАЗС — Автоматизированная контейнерная АЗС', 2)
    doc.add_paragraph('Полностью автоматическая станция без оператора')

    add_table(doc,
        ['Функция', 'Описание'],
        [
            ['Полная автономность', 'Станция работает без персонала 24/7'],
            ['Контроль резервуаров', 'Интеграция с уровнемерами'],
            ['Удалённое управление', 'Все данные в TradeFrame Builder'],
            ['Автоматические уведомления', 'Оповещения о низком уровне топлива'],
            ['Экономичность', 'Минимальные затраты на эксплуатацию'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Идеально подходит для: ').bold = True
    p.add_run('удалённых локаций (карьеры, стройплощадки), корпоративных заправок, '
              'небольших населённых пунктов, придорожных станций.')

    add_heading(doc, '3.2 Дополнительный терминал на стационарной АЗС', 2)
    doc.add_paragraph('Терминал работает параллельно с оператором')

    add_table(doc,
        ['Время', 'Режим работы', 'Описание'],
        [
            ['День', 'Параллельно с оператором', 'Ускорение обслуживания, 2 точки продаж'],
            ['Ночь', 'Самостоятельно', 'Станция работает без персонала'],
            ['Праздники', 'Самостоятельно', 'Не нужно выводить оператора'],
        ]
    )

    add_heading(doc, '3.3 Возможности удалённого управления', 2)
    add_table(doc,
        ['Функция', 'Описание'],
        [
            ['Управление питанием', 'Удаленная перезагрузка станции и терминала'],
            ['Ценообразование', 'Централизованная установка цен на топливо'],
            ['Оповещения', 'Настройка уведомлений (Telegram, Email, SMS)'],
            ['Мониторинг здоровья', 'Контроль работоспособности всех узлов'],
            ['Автокалибровка', 'Калибровка резервуаров на основе данных проливов'],
            ['Анализ запасов', 'Сверка книжных и фактических остатков'],
            ['Отчетность', 'Формирование сменных отчетов'],
        ]
    )

    doc.add_page_break()

    # ===== РАЗДЕЛ 4: ИНСТРУКЦИЯ ПО ИСПОЛЬЗОВАНИЮ =====
    add_heading(doc, '4. Инструкция по использованию терминала', 1)

    doc.add_paragraph(
        'Добро пожаловать! Этот терминал позволяет быстро и удобно оплатить топливо. '
        'Пожалуйста, следуйте этим простым шагам.'
    )

    add_heading(doc, 'Навигация по меню', 2)
    nav_items = [
        ('Далее', 'переход к следующему шагу'),
        ('Назад', 'возврат на предыдущий экран'),
        ('Автоматическая отмена', 'если не нажимать 5 минут, заказ отменится'),
    ]
    for item, desc in nav_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item).bold = True
        p.add_run(f' — {desc}')

    # Способ 1
    add_heading(doc, 'Способ 1: Оплата банковской картой', 2)
    doc.add_paragraph('Это самый простой и быстрый способ заправки.').italic = True

    steps_card = [
        'Снимите пистолет с нужным видом топлива',
        'На главном экране нажмите «Карта банка»',
        'Выберите тип заказа: «На сумму» или «На литры»',
        'Введите сумму или количество литров, нажмите «Далее»',
        'Выберите сторону ТРК, проверьте информацию, нажмите «Далее»',
        'Проверьте итоговую информацию, нажмите «Далее»',
        'Приложите карту к терминалу банка',
        'Заберите чек и заправляйтесь!',
    ]
    for i, step in enumerate(steps_card, 1):
        doc.add_paragraph(f'{i}. {step}')

    p = doc.add_paragraph()
    p.add_run('ВАЖНО! ').bold = True
    p.add_run('Если в бак поместится меньше топлива, разница будет возвращена на карту.')

    # Способ 2
    add_heading(doc, 'Способ 2: Оплата наличными', 2)

    steps_cash = [
        'Снимите пистолет с нужным видом топлива',
        'Внесите купюры в купюроприёмник, нажмите «Далее»',
        'Выберите сторону ТРК, нажмите «Далее»',
        'Подтвердите заказ, нажмите «Далее»',
        'Заберите чек и заправляйтесь!',
    ]
    for i, step in enumerate(steps_cash, 1):
        doc.add_paragraph(f'{i}. {step}')

    p = doc.add_paragraph()
    p.add_run('ВАЖНО! ').bold = True
    p.add_run('Сохраняйте чек! Если заправили не весь объем, на чеке будет номер купона на остаток.')

    # Способ 3
    add_heading(doc, 'Способ 3: Использование купона', 2)

    steps_coupon = [
        'Снимите пистолет с нужным видом топлива',
        'Нажмите «Наличные и купоны», затем «Купон»',
        'Введите номер купона с чека',
        'Выберите сторону ТРК, нажмите «Далее»',
        'Подтвердите заказ, заберите чек и заправляйтесь!',
    ]
    for i, step in enumerate(steps_coupon, 1):
        doc.add_paragraph(f'{i}. {step}')

    # Способ 4
    add_heading(doc, 'Способ 4: Оплата топливной картой', 2)

    steps_fuel_card = [
        'Снимите пистолет с нужным видом топлива',
        'На главном экране нажмите «Топливная карта»',
        'Вставьте карту в картридер',
        'Выберите операцию: Оплата, Дисконт, Баланс или След. кошелёк',
        'Выберите тип заказа: «На сумму» или «На литры»',
        'Введите сумму или литры, нажмите «Далее»',
        'Выберите сторону ТРК, нажмите «Далее»',
        'Подтвердите заказ, заберите чек и заправляйтесь!',
    ]
    for i, step in enumerate(steps_fuel_card, 1):
        doc.add_paragraph(f'{i}. {step}')

    # Способ 5
    add_heading(doc, 'Способ 5: Оплата через онлайн-заказ', 2)

    doc.add_paragraph('Предварительный этап (в мобильном приложении):').bold = True
    steps_online = [
        'Откройте мобильное приложение Агрегатора',
        'Выберите нужную АЗС на карте',
        'Укажите тип топлива и количество',
        'Оплатите заказ онлайн',
        'Снимите пистолет и выберите сторону ТРК',
        'Заправляйтесь! Чек будет в приложении.',
    ]
    for i, step in enumerate(steps_online, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_page_break()

    # ===== РАЗДЕЛ 5: ЭКРАНЫ ТЕРМИНАЛА =====
    add_heading(doc, '5. Экраны терминала', 1)

    screens = [
        ('5.1 Главный экран — Выбор типа оплаты', '1 - экран оплаты.jpg',
         'Три кнопки: Наличные, Топливная карта, Карта банка'),

        ('5.2 Внесение наличных', '2 - внос налички.jpg',
         'Отображается текущая сумма внесённых денег. Кнопки: Отмена, Далее'),

        ('5.3 Меню банковской карты', '3 - оплата МПС.jpg',
         'Варианты: Оплата, Отпуск, Возврат'),

        ('5.4 Меню топливной карты (сценарий 1)', '4.1 - топливная карта (сценарий 1).jpg',
         'Варианты: Оплата, Дисконт, Доп. операции, След. кошелёк'),

        ('5.4 Меню топливной карты (сценарий 2)', '4.2 - топливная карта (сценарий 2).jpg',
         'Альтернативный вариант меню топливной карты'),

        ('5.5 Выбор типа заказа', '5 - выбор типа заказа.jpg',
         'Кнопки: На сумму, На литры'),

        ('5.6 Ввод литров', '6.1 - ввод литров.jpg',
         'Цифровая клавиатура для ввода количества литров'),

        ('5.6 Ввод суммы', '6.2 - ввод суммы.jpg',
         'Цифровая клавиатура для ввода суммы'),

        ('5.7 Выбор колонки', '6.3 - ввод номера колонки.jpg',
         'Клиент вводит номер колонки'),

        ('5.8 Инструкция "Снимите пистолет"', '7.1 - снять пистолет МПС.jpg',
         'Терминал просит снять заправочный пистолет'),

        ('5.9 Подтверждение заказа', '7.2 - итого МПС.jpg',
         'Итоговая информация: тип топлива и количество'),

        ('5.10 Завершение операции', '9 - благодарность.jpg',
         'Экран "Спасибо за покупку!" с напоминанием забрать чек'),
    ]

    for title, image_file, description in screens:
        add_heading(doc, title, 2)

        image_path = os.path.join(IMAGES_DIR, image_file)
        add_image(doc, image_path, width=Inches(4.5))

        doc.add_paragraph()
        doc.add_paragraph(description)
        doc.add_paragraph()

    # Сохранение документа
    doc.save(OUTPUT_FILE)
    print(f"Документ создан: {OUTPUT_FILE}")

if __name__ == "__main__":
    create_document()
